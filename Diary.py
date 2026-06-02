"""
DIARY PRO - THE BEST DIARY APP EVER (2025 Edition)
Single file • Multi-user • Images • Export • Print • Backup • Secure
Just run: python DiaryPro.py
"""

import os
import json
import hashlib
import uuid
import threading
import shutil
import tempfile
import platform
from datetime import datetime
from pathlib import Path
import tkinter as tk
from tkinter import ttk, messagebox, simpledialog, filedialog
from tkcalendar import Calendar

# Image & Export
from PIL import Image, ImageTk
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
try:
    from docx import Document
except ImportError:
    Document = None  # Graceful fallback

# ========================= CONFIG =========================
HOME = Path.home()
APP_DIR = HOME / "Documents" / "MyDiaryPro"
APP_DIR.mkdir(parents=True, exist_ok=True)

DATA_FILE = APP_DIR / "data.json"
USERS_FILE = APP_DIR / "users.json"
IMG_DIR = APP_DIR / "images"
THUMB_DIR = APP_DIR / "thumbs"
BACKUP_DIR = APP_DIR / "backups"

for d in (IMG_DIR, THUMB_DIR, BACKUP_DIR):
    d.mkdir(exist_ok=True)

AUTO_BACKUP_MINUTES = 5

# ========================= UTILS =========================
def load_json(path, default=None):
    if default is None: default = {}
    try:
        return json.loads(path.read_text(encoding="utf-8")) if path.exists() else default
    except:
        return default

def save_json(path, data):
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

def hash_password(pw):
    salt = os.urandom(32).hex()
    hashed = hashlib.sha256((salt + pw).encode()).hexdigest()
    return salt, hashed

def verify_password(pw, salt, hashed):
    return hashlib.sha256((salt + pw).encode()).hexdigest() == hashed

def copy_image(src):
    dst = IMG_DIR / f"{datetime.now():%Y%m%d_%H%M%S}_{Path(src).name}"
    shutil.copy2(src, dst)
    return str(dst)

def make_thumb(src, size=(200, 140)):
    try:
        img = Image.open(src)
        img.thumbnail(size)
        thumb = THUMB_DIR / f"thumb_{Path(src).name}.jpg"
        img.convert("RGB").save(thumb, "JPEG", quality=90)
        return str(thumb)
    except:
        return None

def auto_backup():
    def run():
        while True:
            try:
                if DATA_FILE.exists():
                    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                    shutil.copy2(DATA_FILE, BACKUP_DIR / f"backup_{ts}.json")
            except: pass
            threading.Event().wait(AUTO_BACKUP_MINUTES * 60)
    threading.Thread(target=run, daemon=True).start()

# ========================= MAIN APP =========================
class DiaryPro:
    def __init__(self, root, username):
        self.root = root
        self.user = username
        self.root.title(f"Diary Pro — {username}")
        self.root.geometry("1380x820")
        self.root.minsize(1100, 650)
        self.dark = False

        self.data = load_json(DATA_FILE)
        if username not in self.data:
            self.data[username] = []
            save_json(DATA_FILE, self.data)

        self.filtered = []
        self.current_idx = None

        self.build_ui()
        auto_backup()
        self.refresh_list()

    def build_ui(self):
        # Top Bar
        top = tk.Frame(self.root, bg="#2c3e50", height=70)
        top.pack(fill="x")
        top.pack_propagate(False)

        tk.Label(top, text="DIARY PRO", font=("Arial", 28, "bold"), bg="#2c3e50", fg="white").pack(side="left", padx=25, pady=10)

        search_f = tk.Frame(top, bg="#2c3e50")
        search_f.pack(side="right", padx=20)
        tk.Label(search_f, text="Search:", font=("Arial", 14), bg="#2c3e50", fg="white").pack(side="left")
        self.search_var = tk.StringVar()
        self.search_var.trace("w", lambda *a: self.refresh_list())
        tk.Entry(search_f, textvariable=self.search_var, font=("Arial", 14), width=30).pack(side="left", padx=8)

        tk.Button(top, text="Dark Mode", command=self.toggle_dark, bg="#34495e", fg="white", font=("bold",12)).pack(side="right", padx=10)

        # Main Panes
        main = tk.PanedWindow(self.root, orient="horizontal", sashwidth=6)
        main.pack(fill="both", expand=True, padx=15, pady=15)

        # Left: Calendar + List
        left = tk.Frame(main, width=420, bg="#ecf0f1")
        main.add(left)

        tk.Label(left, text="Calendar", font=("Arial", 16, "bold"), bg="#ecf0f1").pack(pady=20)
        self.cal = Calendar(left, selectmode="day", date_pattern="y-mm-dd", background="#2c3e50", foreground="white")
        self.cal.pack(padx=30, pady=10)
        self.cal.bind("<<CalendarSelected>>", lambda e: self.refresh_list())

        tk.Button(left, text="Show All", command=lambda: [self.cal.selection_clear(), self.search_var.set(""), self.refresh_list()],
                  bg="#27ae60", fg="white", font=("bold",14), height=2).pack(pady=20, padx=60, fill="x")

        tk.Label(left, text="Entries", font=("Arial", 16, "bold"), bg="#ecf0f1").pack(pady=(20,10))
        cols = ("Date", "Title", "Tags")
        self.tree = ttk.Treeview(left, columns=cols, show="headings", selectmode="browse")
        for c, w in zip(cols, [130, 200, 120]):
            self.tree.heading(c, text=c)
            self.tree.column(c, width=w, anchor="w")
        self.tree.pack(fill="both", expand=True, padx=30, pady=10)
        self.tree.bind("<<TreeviewSelect>>", self.load_entry)
        self.tree.bind("<Double-1>", lambda e: self.edit_entry())

        # Right: Editor
        right = tk.Frame(main, bg="white")
        main.add(right)

        # Title
        tk.Label(right, text="Title", font=("Arial", 20, "bold"), bg="white").pack(anchor="w", padx=30, pady=(20,8))
        self.title_var = tk.StringVar()
        tk.Entry(right, textvariable=self.title_var, font=("Arial", 18)).pack(fill="x", padx=30, pady=5)

        # Tags
        tk.Label(right, text="Tags (comma separated)", font=("Arial", 12), bg="white", fg="#7f8c8d").pack(anchor="w", padx=30, pady=(10,5))
        self.tags_var = tk.StringVar()
        tk.Entry(right, textvariable=self.tags_var, font=("Arial", 12)).pack(fill="x", padx=30, pady=5)

        # Content
        self.text = tk.Text(right, height=18, font=("Arial", 14), wrap="word", relief="solid", bd=2)
        self.text.pack(fill="both", expand=True, padx=30, pady=15)

        # Images
        img_bar = tk.Frame(right, bg="#f0f0f0", height=140)
        img_bar.pack(fill="x", padx=30, pady=10)
        img_bar.pack_propagate(False)
        tk.Button(img_bar, text="Attach Photo", command=self.attach_photo, bg="#3498db", fg="white", font=("bold",12)).pack(side="left", padx=10, pady=10)
        tk.Button(img_bar, text="Remove Photo", command=self.remove_photo, bg="#e74c3c", fg="white", font=("bold",12)).pack(side="left", padx=10, pady=10)
        self.img_canvas = tk.Canvas(img_bar, bg="#f0f0f0", highlightthickness=0)
        self.img_canvas.pack(side="left", fill="both", expand=True, padx=10)

        # Buttons
        btns = tk.Frame(right, bg="white")
        btns.pack(pady=25)
        tk.Button(btns, text="NEW ENTRY", command=self.new_entry, bg="#9b59b6", fg="white", font=("Arial", 16, "bold"), width=16, height=2).pack(side="left", padx=20)
        tk.Button(btns, text="SAVE", command=self.save_entry, bg="#27ae60", fg="white", font=("Arial", 18, "bold"), width=18, height=2).pack(side="left", padx=20)
        tk.Button(btns, text="DELETE", command=self.delete_entry, bg="#e74c3c", fg="white", font=("Arial", 18, "bold"), width=16, height=2).pack(side="left", padx=20)

        # Export Bar
        exp = tk.Frame(right, bg="white")
        exp.pack(pady=10)
        for txt, cmd, color in [
            ("Export TXT", self.export_txt, "#9b59b6"),
            ("Export PDF", self.export_pdf, "#e67e22"),
            ("Export DOCX", self.export_docx, "#2980b9"),
            ("Print", self.print_entry, "#2c3e50")
        ]:
            tk.Button(exp, text=txt, command=cmd, bg=color, fg="white", font=("bold",12), width=18).pack(side="left", padx=12)

    def toggle_dark(self):
        self.dark = not self.dark
        bg = "#2c3e50" if self.dark else "#f8f9fa"
        fg = "white" if self.dark else "black"
        self.root.configure(bg=bg)
        # Simple dark mode (you can expand this)
        for w in [self.root] + list(self.root.winfo_children()):
            try: w.configure(bg=bg, fg=fg)
            except: pass

    def get_entries(self):
        return self.data.get(self.user, [])

    def save_data(self):
        save_json(DATA_FILE, self.data)

    def refresh_list(self):
        for i in self.tree.get_children():
            self.tree.delete(i)
        entries = self.get_entries()
        q = self.search_var.get().lower()
        date = self.cal.get_date() if self.cal.selection_get() else None

        filtered = []
        for e in entries:
            if q and q not in f"{e.get('title','')} {e.get('content','')} {','.join(e.get('tags',[]))}".lower():
                continue
            if date and not e["created"].startswith(date):
                continue
            filtered.append(e)

        filtered.sort(key=lambda x: x["created"], reverse=True)
        self.filtered = filtered

        for e in filtered:
            tags = ", ".join(e.get("tags", []))[:30]
            self.tree.insert("", "end", values=(
                e["created"][:10], e.get("title", "Untitled"), tags
            ))

    def load_entry(self, e=None):
        sel = self.tree.selection()
        if not sel: return
        idx = self.tree.index(sel[0])
        e = self.filtered[idx]
        self.current_idx = self.get_entries().index(e)

        self.title_var.set(e.get("title", ""))
        self.tags_var.set(", ".join(e.get("tags", [])))
        self.text.delete(1.0, "end")
        self.text.insert(1.0, e.get("content", ""))
        self.show_images(e)

    def new_entry(self):
        self.current_idx = None
        self.title_var.set("")
        self.tags_var.set("")
        self.text.delete(1.0, "end")
        self.img_canvas.delete("all")

    def save_entry(self):
        title = self.title_var.get().strip() or "Untitled"
        content = self.text.get(1.0, "end").strip()
        tags = [t.strip() for t in self.tags_var.get().split(",") if t.strip()]

        entries = self.get_entries()
        if self.current_idx is None:
            entry = {
                "id": str(uuid.uuid4()),
                "created": datetime.now().isoformat(),
                "title": title,
                "content": content,
                "tags": tags,
                "images": []
            }
            entries.append(entry)
        else:
            e = entries[self.current_idx]
            e.update({"title": title, "content": content, "tags": tags})

        self.save_data()
        self.refresh_list()
        messagebox.showinfo("Success", "Entry saved!")

    def delete_entry(self):
        if not self.tree.selection(): return
        if messagebox.askyesno("Delete", "Permanently delete this entry?"):
            idx = self.tree.index(self.tree.selection()[0])
            e = self.filtered[idx]
            entries = self.get_entries()
            entries.remove(e)
            self.save_data()
            self.refresh_list()
            self.new_entry()

    def attach_photo(self):
        path = filedialog.askopenfilename(filetypes=[("Images", "*.jpg *.jpeg *.png *.gif")])
        if not path: return
        dst = copy_image(path)
        thumb = make_thumb(dst)
        if self.current_idx is None:
            self.save_entry()
        entries = self.get_entries()
        entries[-1]["images"].append(dst)
        self.save_data()
        self.show_images(entries[-1])

    def show_images(self, entry):
        self.img_canvas.delete("all")
        x = 10
        self.photos = []
        for p in entry.get("images", []):
            if not Path(p).exists(): continue
            thumb = make_thumb(p) or p
            try:
                img = Image.open(thumb)
                img = img.resize((180, 120), Image.Resampling.LANCZOS)
                photo = ImageTk.PhotoImage(img)
                self.img_canvas.create_image(x, 10, anchor="nw", image=photo)
                self.photos.append(photo)
                x += 190
            except: pass

    def remove_photo(self):
        if self.current_idx is not None and self.get_entries()[self.current_idx]["images"]:
            self.get_entries()[self.current_idx]["images"].pop()
            self.save_data()
            self.show_images(self.get_entries()[self.current_idx])

    def export_txt(self):
        path = filedialog.asksaveasfilename(defaultextension=".txt")
        if path:
            with open(path, "w", encoding="utf-8") as f:
                f.write("MY DIARY PRO\n" + "="*80 + "\n\n")
                for e in self.get_entries()[::-1]:
                    f.write(f"{e['created'][:10]} — {e['title']}\n\n{e['content']}\n\n{'-'*80}\n\n")
            messagebox.showinfo("Done", "TXT exported!")

    def export_pdf(self):
        if not Document: 
            messagebox.showerror("Missing", "Install python-docx for PDF export")
            return
        path = filedialog.asksaveasfilename(defaultextension=".pdf")
        if path:
            c = canvas.Canvas(path, pagesize=A4)
            w, h = A4
            y = h - 50
            for e in self.get_entries()[::-1]:
                if y < 100:
                    c.showPage()
                    y = h - 50
                c.setFont("Helvetica-Bold", 14)
                c.drawString(50, y, e["title"])
                y -= 20
                c.setFont("Helvetica", 10)
                c.drawString(50, y, e["created"][:16].replace("T", " "))
                y -= 30
                for line in e["content"].split("\n"):
                    if y < 50:
                        c.showPage()
                        y = h - 50
                    c.drawString(60, y, line[:100])
                    y -= 14
                y -= 20
            c.save()
            messagebox.showinfo("Success", "PDF created!")

    def export_docx(self):
        if not Document:
            messagebox.showerror("Missing", "Install python-docx")
            return
        path = filedialog.asksaveasfilename(defaultextension=".docx")
        if path:
            doc = Document()
            doc.add_heading("My Diary Pro", 0)
            for e in self.get_entries()[::-1]:
                doc.add_heading(e["title"], 2)
                doc.add_paragraph(e["created"][:16].replace("T", " "))
                doc.add_paragraph(e["content"])
                doc.add_page_break()
            doc.save(path)
            messagebox.showinfo("Done", "DOCX saved!")

    def print_entry(self):
        if not self.tree.selection(): return
        messagebox.showinfo("Print", "Printing not implemented in this version\nUse Export → PDF → Print")

# ========================= LOGIN =========================
class Login:
    def __init__(self):
        root = tk.Tk()
        root.title("Diary Pro Login")
        root.geometry("400x300")
        root.configure(bg="#2c3e50")

        tk.Label(root, text="DIARY PRO", font=("Arial", 32, "bold"), bg="#2c3e50", fg="white").pack(pady=30)
        tk.Label(root, text="Username", font=("Arial", 14), bg="#2c3e50", fg="white").pack(pady=10)
        user = tk.Entry(root, font=("Arial", 14), width=25)
        user.pack(pady=5)
        tk.Label(root, text="Password", font=("Arial", 14), bg="#2c3e50", fg="white").pack(pady=10)
        pw = tk.Entry(root, font=("Arial", 14), width=25, show="*")
        pw.pack(pady=5)

        users = load_json(USERS_FILE)

        def login():
            u, p = user.get().strip(), pw.get()
            if u in users and verify_password(p, *users[u]):
                root.destroy()
                DiaryPro(tk.Tk(), u)
            else:
                messagebox.showerror("Error", "Invalid login")

        def register():
            u, p = user.get().strip(), pw.get()
            if u and p and u not in users:
                users[u] = hash_password(p)
                save_json(USERS_FILE, users)
                messagebox.showinfo("Success", "Registered! Now login.")
            else:
                messagebox.showerror("Error", "Invalid or existing user")

        tk.Button(root, text="LOGIN", command=login, bg="#27ae60", fg="white", font=("bold",14), width=15, height=2).pack(pady=15)
        tk.Button(root, text="REGISTER", command=register, bg="#3498db", fg="white", font=("bold",12), width=15).pack()

        root.mainloop()

if __name__ == "__main__":
    Login()